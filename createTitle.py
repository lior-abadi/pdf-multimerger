from docx2pdf import convert
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil

class FolderHeader:
    def createHeader(self, name, dir):  
        document = Document()
        p = document.add_paragraph()
        name = str(name).upper()
        run = p.add_run(f'{name}')
        paragraph_format = p.paragraph_format

        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        font = run.font
        font.size = Pt(35)
        font.name = 'Calibri'
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)

        document.save(f'./{dir}/00. Aa {name}.docx')
        convert(f'./{dir}/00. Aa {name}.docx', f'./{dir}/00. Aa {name}.pdf')
        os.remove(f'./{dir}/00. Aa {name}.docx')
    
    def addBlank(self, dir):  
        shutil.copy(r"./ZZ_blank_page.pdf", dir)
  
            
